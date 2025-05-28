import hashlib
from docx import Document
import pdfkit
import argparse
import os
import json
import time
import re
import pandas as pd
import subprocess
from jinja2 import Template
import pandas as pd
import requests


from mon import cupons_aplicador
from notify import wht_send
from datetime import datetime
from collections import defaultdict
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import NoSuchElementException, StaleElementReferenceException, TimeoutException
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from tabulate import tabulate
from common import log, env
from openpyxl.styles import Alignment


PAGINATION_LABEL = "ml_seguinte"
WAIT_TIME = 15
BASE_DIR = 'snapshots'
IMAGENS_DIR = os.path.join(BASE_DIR, 'images')
os.makedirs(IMAGENS_DIR, exist_ok=True)

CLASSIFICACAO = {
    "A1 Mini": {
        "combo_keywords": ["mini combo", "mini ams", "mini multicolor"],
        "sem_combo_keywords": [
            "mini sem ams", 
            "mini sem combo", 
            "pf002", 
            "alta performance"
        ]
    },
    "A1": {
        "combo_keywords": ["combo", "ams", "multicolor"],
        "sem_combo_keywords": [
            "sem ams", 
            "sem combo", 
            "fdm", 
            "pf002", 
            "pf001", 
            "alta performance", 
            "aberta"
        ]
    },
    "P1S": {
        "combo_keywords": ["p1s combo", "p1s ams"],
        "sem_combo_keywords": ["p1s sem ams", "p1s -", "p1s"]
    },
    "P1P": {
        "keywords": ["p1p"]
    },
    "Outros": {
        "exclusoes": ["filamento", "sistema automático", "suporte", "curso", "brinde"]
    }
}

IGNORADOS = [
    {"produto": "Impressora 3d Bambu Lab A1 Mini (sem Ams) Cor Branco 127/220v", "preco": 2.499},
    {"produto": "Ams Sistema Automático De Materiais P1s - X1c", "preco": 4650.00},
    {"produto": "Bambu Lab Ams Sa001 Sistema Automático De Materiais P1s X1c", "preco": 4650.00},
    {"produto": "Ams Sistema Automático De Materiais P1s - X1c", "preco": 4650.00},
    {"produto": "Bambu Lab Ams Sa001 Sistema Automático De Materiais P1s X1c", "preco": 4650.00}
]

def enviar_para_google_drive(caminho_local, pasta_remota='WHTAutomacao'):
    comando = ["rclone", "copy", caminho_local, f"{pasta_remota}:"]
    try:
        subprocess.run(comando, check=True)
        print(f"Arquivo {caminho_local} enviado para Google Drive na pasta {pasta_remota}.")
        link_publico = f"https://drive.google.com/drive/folders/{env.RCLONE_FOLDER_ID}"
        return link_publico
    except subprocess.CalledProcessError as e:
        print(f"Erro ao enviar para o Google Drive: {e}")
        return None

def extrair_codigo_para_link(link):
    match = re.search(r"(MLB\d+)(?!\d)", link)
    if match:
        return f"https://www.mercadolivre.com.br/p/{match.group(1)}"
    return link

def deve_ignorar(titulo, preco):
    for item in IGNORADOS:
        #if item['produto'].lower() == titulo.lower() and abs(item['preco'] - preco) < 0.01:
        if item['produto'].lower() == titulo.lower():
            return True
    return False

def baixar_imagem(url, pasta_destino=IMAGENS_DIR):
    os.makedirs(pasta_destino, exist_ok=True)
    try:
        response = requests.get(url, stream=True, timeout=10)
        response.raise_for_status()
        nome_arquivo = hashlib.md5(url.encode()).hexdigest() + '.jpg'
        caminho_arquivo = os.path.join(pasta_destino, nome_arquivo)
        with open(caminho_arquivo, 'wb') as out_file:
            for chunk in response.iter_content(chunk_size=8192):
                out_file.write(chunk)
        return os.path.relpath(caminho_arquivo, BASE_DIR)
    except Exception as e:
        print(f"[ERRO] Não foi possível baixar imagem {url}: {e}")
        return None

def exibir_resumo_por_classificacao(produtos):
    resumo = defaultdict(list)

    for produto, info in produtos.items():
        resumo[info['classificacao']].append({
            'preco': info['preco'],
            'link_reduzido': info['link_reduzido']
        })

    tabela_resumo = []
    for classe, itens in sorted(resumo.items()):
        menor_item = min(itens, key=lambda x: x['preco'])
        menor_preco = menor_item['preco']
        link = menor_item['link_reduzido']

        tabela_resumo.append([classe, len(itens), f"R$ {menor_preco:.2f}", link])

    print("\nResumo por classificação:")
    print(tabulate(tabela_resumo, headers=['Classificação', 'Quantidade', 'Menor Preço', 'Link'], tablefmt='grid'))

def enviar_resumo_whatsapp(produtos):
    resumo = defaultdict(list)
    for produto, info in produtos.items():
        resumo[info['classificacao']].append({
            'preco': info['preco'],
            'link_reduzido': info['link_reduzido']
        })

    for classe, itens in sorted(resumo.items()):
        menor_item = min(itens, key=lambda x: x['preco'])
        menor_preco = menor_item['preco']
        link = menor_item['link_reduzido']
        message = f"{classe}: {len(itens)} itens\nMenor: R$ {menor_preco:.2f}\n{link}"
        wht_send.send_whatsapp_message(env.PHONE, env.API_KEY, message)
        time.sleep(2)

def enviar_link_whatsapp(link_arquivo):
    message = f"Arquivo disponível: {link_arquivo}"
    wht_send.send_whatsapp_message(env.PHONE, env.API_KEY, message)

def setup_driver(silent=False):
    options = Options()
    options.add_argument(r"--user-data-dir=C:\\tools\\SeleniumProfile")
    options.add_argument("--start-maximized")
    options.add_argument("--disable-blink-features=AutomationControlled")

    if silent:
        options.add_argument('--headless')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')

    return webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

def find_seguinte_button(driver):
    try:
        link = driver.find_element(By.XPATH, "//span[text()='Seguinte']/ancestor::a")
        return link if link.is_displayed() and link.is_enabled() else None
    except (NoSuchElementException, StaleElementReferenceException):
        return None

def click_next_page(driver):
    for tentativa in range(1, 4):
        link = find_seguinte_button(driver)
        if not link:
            log_and_print("🛑 Botão 'Seguinte' não encontrado ou não visível/enabled. Fim da execução.", PAGINATION_LABEL)
            return False

        try:
            driver.execute_script("arguments[0].scrollIntoView(true);", link)
            time.sleep(1)
            log_and_print(f"➡️ Clicando em 'Seguinte'... (tentativa {tentativa})", PAGINATION_LABEL)
            driver.execute_script("arguments[0].click();", link)
            WebDriverWait(driver, WAIT_TIME).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, ".ui-search-layout__item"))
            )
            return True

        except (StaleElementReferenceException, TimeoutException, NoSuchElementException) as e:
            log_and_print(f"⚠️ Problema na tentativa {tentativa}: {e}. Aguardando e tentando novamente...", PAGINATION_LABEL)
            time.sleep(2)

    log_and_print("❌ Não foi possível clicar em 'Seguinte' após múltiplas tentativas.", PAGINATION_LABEL)
    return False

def log_and_print(message, label='web_mon'):
    print(message)
    log.write_log(label, message)

def wait_for_products(driver):
    try:
        WebDriverWait(driver, WAIT_TIME).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, ".ui-search-layout__item"))
        )
        return True
    except TimeoutException:
        log_and_print("⚠️ Timeout ao esperar itens na página. Fim da coleta.")
        return False

def classificar_produto(titulo):
    titulo_lower = titulo.lower()

    # ✅ Prioridade A1 Mini
    if 'mini' in titulo_lower:
        regras = CLASSIFICACAO.get("A1 Mini", {})
        if any(kw in titulo_lower for kw in regras.get("sem_combo_keywords", [])):
            return "A1 Mini Sem Combo"
        if any(kw in titulo_lower for kw in regras.get("combo_keywords", [])):
            return "A1 Mini Combo"
        return "A1 Mini"

    # ✅ Demais categorias
    for categoria in ["P1P", "P1S", "A1"]:
        regras = CLASSIFICACAO.get(categoria, {})

        # ✅ SEM COMBO sempre primeiro!
        if any(kw in titulo_lower for kw in regras.get("sem_combo_keywords", [])):
            return f"{categoria} Sem Combo"

        # ✅ Depois combo
        if any(kw in titulo_lower for kw in regras.get("combo_keywords", [])):
            return f"{categoria} Combo"

        # ✅ Por fim, apenas o nome
        if any(kw in titulo_lower for kw in regras.get("keywords", [])):
            return categoria

    # ✅ Caso ainda não classificou e contém "A1" puro
    if "a1" in titulo_lower:
        # Se tem "sem ams" ou outros, Sem Combo
        if any(kw in titulo_lower for kw in CLASSIFICACAO["A1"].get("sem_combo_keywords", [])):
            return "A1 Sem Combo"
        # Se não tem nada, default para "A1"
        return "A1"

    # ✅ Por fim, Exclusões
    for ex in CLASSIFICACAO["Outros"].get("exclusoes", []):
        if ex in titulo_lower:
            return "Outros"

    return "Não Classificado"

def coletar_produtos(driver, url, uma_pagina=False):
    consolidado = {}
    paginas = 0
    driver.get(url)

    while True:
        time.sleep(3)
        if not wait_for_products(driver):
            break

        parse_products(driver, consolidado)
        paginas += 1
        log_and_print(f"[INFO] Total parcial de produtos coletados: {len(consolidado)} após {paginas} página(s).")

        if uma_pagina:
            log_and_print("[INFO] Parando após primeira página, conforme solicitado.")
            break

        if not find_seguinte_button(driver):
            log_and_print("[INFO] Não há mais páginas. Fim da coleta.")
            break

        if not click_next_page(driver):
            break

    log_and_print(f"[INFO] Total produtos coletados: {len(consolidado)} em {paginas} página(s).")
    return consolidado, paginas

def exibir_tabela(produtos, paginas):
    tabela = [
        [produto, f"R$ {info['preco']:.2f}", info['classificacao'], 'Sim' if info['existe_cupom'] else 'Não']
        for produto, info in produtos.items()
    ]
    print(tabulate(tabela, headers=['Produto', 'Menor Preço', 'Classificação', 'Existe Cupom'], tablefmt='grid'))
    print(f"\nTotal de páginas analisadas: {paginas}")
    print(f"Total de produtos coletados: {len(produtos)}")

def parse_products(driver, consolidado):
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(driver.page_source, 'html.parser')
    for item in soup.select('.ui-search-layout__item'):
        titulo_tag = item.select_one('.poly-component__title-wrapper .poly-component__title')
        preco_tag = item.select_one('.poly-price__current .andes-money-amount__fraction')
        link_tag = item.select_one('a')
        parcela_tag = item.select_one('.poly-price__installments')
        desconto_tag = item.select_one('.andes-money-amount__discount')

        if not titulo_tag or not preco_tag or not link_tag:
            continue

        titulo = titulo_tag.get_text(strip=True)
        preco_str = preco_tag.get_text(strip=True).replace('.', '').replace(',', '.')
        link = link_tag.get('href')

        possui_18x = False
        valor_parcela = "-"

        if parcela_tag:
            parcela_text = parcela_tag.get_text(strip=True).lower()
            possui_18x = '18x' in parcela_text

            # 🔴 NOVO: buscar o valor da parcela dentro da hierarquia correta
            valor_parcela_tag = parcela_tag.select_one('.andes-money-amount .andes-money-amount__fraction')
            if valor_parcela_tag:
                valor_parcela_str = valor_parcela_tag.get_text(strip=True)
                valor_parcela = f"R$ {valor_parcela_str}"

        existe_cupom = False
        if desconto_tag and '% OFF' in desconto_tag.get_text(strip=True):
            existe_cupom = True

        if not possui_18x:
            continue

        try:
            preco = float(preco_str)
            if deve_ignorar(titulo, preco):
                continue
            if titulo not in consolidado or preco < consolidado[titulo]['preco']:
                consolidado[titulo] = {
                    "preco": preco,
                    "classificacao": classificar_produto(titulo),
                    "possui_18x": possui_18x,
                    "valor_parcela": valor_parcela,
                    "link": link,
                    "link_reduzido": extrair_codigo_para_link(link)
                }
        except ValueError:
            continue

def salvar_xlsx(produtos, nome_arquivo='resultado'):
    from openpyxl.utils import get_column_letter
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, PatternFill

    pasta = os.path.join(BASE_DIR, 'ml_scraper_v2')
    os.makedirs(pasta, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_completo = f"{nome_arquivo}_{timestamp}.xlsx"
    caminho_local = os.path.join(pasta, nome_completo)

    dados = [{
        "Produto": produto,
        "Menor Preço": info['preco'],
        "Classificação": info['classificacao'],
        "Possui 18x": info['possui_18x'],
        "Valor Parcela": info['valor_parcela'],
        "Existe Cupom": info['existe_cupom'],
        "Link": info['link']
    } for produto, info in produtos.items()]

    df = pd.DataFrame(dados).sort_values(by=["Produto", "Menor Preço"])
    df.to_excel(caminho_local, index=False, engine='openpyxl')

    wb = load_workbook(caminho_local)
    ws = wb.active

    for col in ws.columns:
        max_length = max(len(str(cell.value)) for cell in col if cell.value)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_length + 2, 15)

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical='top')
            cell.fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    wb.save(caminho_local)
    print(f"Arquivo XLSX salvo em: {caminho_local}")
    return caminho_local

def salvar_docx(produtos, nome_arquivo='resultado'):
    os.makedirs(os.path.join(BASE_DIR, 'ml_scraper_v2'), exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_completo = f"{nome_arquivo}_{timestamp}.docx"
    caminho_local = os.path.join(os.path.join(BASE_DIR, 'ml_scraper_v2'), nome_completo)

    doc = Document()
    doc.add_heading('Resumo por Classificação', level=1)

    resumo = defaultdict(list)
    for produto, info in produtos.items():
        resumo[info['classificacao']].append(info['preco'])

    for classe, precos in sorted(resumo.items()):
        menor_preco = min(precos)
        doc.add_paragraph(f"{classe}: {len(precos)} itens | Menor: R$ {menor_preco:.2f}")

    doc.add_heading('Tabela Completa', level=2)

    for produto, info in produtos.items():
        doc.add_paragraph(
            f"{produto} | R$ {info['preco']:.2f} | {info['classificacao']} | 18x: {'Sim' if info['possui_18x'] else 'Não'} | Cupom: {'Sim' if info['existe_cupom'] else 'Não'} | Link: {info['link']}"
        )

    doc.save(caminho_local)
    print(f"Arquivo DOCX salvo em: {caminho_local}")
    return caminho_local

def salvar_html(produtos, nome_arquivo='resultado'):
    import os
    from datetime import datetime
    from jinja2 import Template
    import pandas as pd

    pasta = os.path.join(BASE_DIR, 'ml_scraper_v2')
    os.makedirs(pasta, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_completo = f"{nome_arquivo}_{timestamp}.html"
    caminho_local = os.path.join(pasta, nome_completo)

    # Agrupar por classificação
    agrupado = defaultdict(list)
    for produto, info in produtos.items():
        agrupado[info.get('classificacao', 'Não Classificado')].append({
            "Produto": produto,
            "Menor Preço": f"R$ {info.get('preco', 0):.2f}",
            "PrecoNum": info.get('preco', 0),  # campo auxiliar numérico
            "Valor Parcela": info.get('valor_parcela', '-'),
            "Possui 18x": 'Sim' if info.get('possui_18x', False) else 'Não',
            "Existe Cupom": 'Sim' if info.get('existe_cupom', False) else 'Não',
            "Link": f'<a href="{info.get("link", "#")}" target="_blank">Ver Produto</a>'
        })

    # Resumo
    resumo_html = "<ul>"
    for classe in sorted(agrupado.keys()):
        precos = [p["PrecoNum"] for p in agrupado[classe]]
        menor_preco = min(precos) if precos else 0
        resumo_html += f"<li><strong>{classe}</strong>: {len(precos)} itens | Menor: R$ {menor_preco:.2f}</li>"
    resumo_html += "</ul>"

    # Criar tabelas por classificação
    tabelas_html = ""
    for classe in sorted(agrupado.keys()):
        df = pd.DataFrame(agrupado[classe])
        # Ordenar por PrecoNum, depois Produto
        df = df.sort_values(by=["PrecoNum", "Produto"])
        df = df.drop(columns=["PrecoNum"])  # remover campo auxiliar
        tabela_html = df.to_html(escape=False, index=False)
        tabelas_html += f"<h2>{classe}</h2>{tabela_html}"

    # Template HTML
    template = Template("""
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Resumo Mercado Livre</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; }
            table { border-collapse: collapse; width: 100%; margin-bottom: 30px; }
            th, td { border: 1px solid #ddd; padding: 8px; }
            th { background-color: #f2f2f2; }
        </style>
    </head>
    <body>
        <h1>Resumo por Classificação</h1>
        {{ resumo }}
        <h1>Tabelas por Classificação</h1>
        {{ tabelas }}
    </body>
    </html>
    """)

    html_content = template.render(resumo=resumo_html, tabelas=tabelas_html)

    with open(caminho_local, 'w', encoding='utf-8') as f:
        f.write(html_content)

    print(f"Arquivo HTML salvo em: {caminho_local}")
    return caminho_local

def salvar_pdf(produtos, nome_arquivo='resultado'):
    caminho_html = salvar_html(produtos, nome_arquivo)
    caminho_pdf = caminho_html.replace('.html', '.pdf')

    try:
        pdfkit.from_file(
            caminho_html,
            caminho_pdf,
            options={'enable-local-file-access': ''}
        )
        print(f"Arquivo PDF salvo em: {caminho_pdf}")
        return caminho_pdf
    except Exception as e:
        print(f"Erro ao gerar PDF: {e}")
        return None

def salvar_json(produtos, nome_arquivo='resultado'):
    pasta = os.path.join(BASE_DIR, 'ml_scraper_v2')
    os.makedirs(pasta, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    caminho = os.path.join(pasta, f"{nome_arquivo}_{timestamp}.json")

    produtos_ordenados = sorted([{
        "produto": produto,
        "menor_preco": info['preco'],
        "classificacao": str(info['classificacao']),
        "possui_18x": info.get('possui_18x', False),
        "valor_parcela": info.get('valor_parcela', '-'),
        "existe_cupom": info.get('existe_cupom', False),
        "link": info.get('link', ''),
        "link_reduzido": info.get('link_reduzido', '')
    } for produto, info in produtos.items()], key=lambda x: (x['produto'], x['menor_preco']))

    with open(caminho, 'w', encoding='utf-8') as f:
        json.dump(produtos_ordenados, f, ensure_ascii=False, indent=4)

    print(f"Arquivo JSON salvo em: {caminho}")
    return caminho

def enviar_resumo_whatsapp_tabela(produtos):
    resumo = defaultdict(list)
    for produto, info in produtos.items():
        resumo[info.get('classificacao', 'Não Classificado')].append({
            'produto': produto,
            'preco': info.get('preco', 0),
            'link_reduzido': info.get('link_reduzido', '#')
        })

    classes_desejadas = {"P1P", "P1S Sem Combo"}

    mensagens = []

    for classe in sorted(resumo.keys()):
        if classe not in classes_desejadas:
            continue
        mensagens.append("--------------")
        itens_ordenados = sorted(resumo[classe], key=lambda x: x['preco'])
        top3 = itens_ordenados[:3]
        for item in top3:
            preco = f"R$ {item['preco']:.2f}"
            link = item['link_reduzido']
            mensagens.append(f"{classe} {preco} {link}")

    mensagem_final = "\n".join(mensagens)

    wht_send.send_whatsapp_message(env.PHONE, env.API_KEY, mensagem_final)
    print("[INFO] Resumo filtrado com top 3 menores preços enviado via WhatsApp.")

def enviar_para_google_drive(caminho_local, pasta_remota='WHTAutomacao'):
    import subprocess
    import os

    nome_arquivo = os.path.basename(caminho_local)
    destino = f"{pasta_remota}:{nome_arquivo}"

    try:
        subprocess.run(["rclone", "copy", caminho_local, f"{pasta_remota}:"], check=True)
        print(f"Arquivo {caminho_local} enviado para Google Drive na pasta {pasta_remota}.")

        # Gerar link público
        resultado = subprocess.run(["rclone", "link", destino], capture_output=True, text=True, check=True)
        link_publico = resultado.stdout.strip()
        print(f"Link público: {link_publico}")
        return link_publico

    except subprocess.CalledProcessError as e:
        print(f"Erro ao enviar ou gerar link: {e}")
        return None

def carregar_json(caminho_json):
    with open(caminho_json, 'r', encoding='utf-8') as f:
        return json.load(f)
    
def comparar_precos_e_alertar(json_atual, json_anterior):
    # Criar dicionário de consulta: {classificacao: [(produto, preco, link)]}
    atual = agrupar_por_classificacao(json_atual)
    anterior = agrupar_por_classificacao(json_anterior)

    for classificacao, lista_atual in atual.items():
        if classificacao not in anterior:
            continue  # Não existia antes, ignora ou alerta de novo produto

        menor_atual = min(lista_atual, key=lambda x: x['menor_preco'])
        menor_anterior = min(anterior[classificacao], key=lambda x: x['menor_preco'])

        if menor_atual['menor_preco'] < menor_anterior['menor_preco']:
            produto = menor_atual['produto']
            preco = menor_atual['menor_preco']
            link = menor_atual.get('link_reduzido', menor_atual['link'])
            mensagem = f"🔻 Queda de preço!\n{classificacao}: {produto}\nNovo: R$ {preco:.2f}\n{link}"
            wht_send.send_whatsapp_message(env.PHONE, env.API_KEY, mensagem)
            print(f"[INFO] Alerta enviado para {classificacao}")

def agrupar_por_classificacao(lista_produtos):
    from collections import defaultdict
    agrupado = defaultdict(list)
    for item in lista_produtos:
        key = item['classificacao']
        if isinstance(key, list):  # ✅ proteção extra
            key = ' '.join(map(str, key))
        agrupado[key].append(item)
    return agrupado

def main():
    parser = argparse.ArgumentParser(description='Mercado Livre Scraper')
    parser.add_argument('--uma_pagina', action='store_true', help='Analisar apenas a primeira página')
    parser.add_argument('--pegar_cupons', action='store_true', help='Aplicar cupons antes da coleta')
    parser.add_argument('--apenas_cupons', action='store_true', help='Executar apenas aplicação de cupons')
    parser.add_argument('--whatsapp', action='store_true', help='Enviar Whatsapp')
    parser.add_argument('--silent', action='store_true', help='Executar em modo silencioso (headless)')
    parser.add_argument('--url', type=str, default='https://lista.mercadolivre.com.br/informatica/novo/_PriceRange_2100-9100_Installments_YES_BRAND_22935733_NoIndex_True', help='URL da página para coleta')
    args = parser.parse_args()

    driver = setup_driver(silent=args.silent)

    try:
        if args.apenas_cupons:
            log.info("run_compara_aviso_multi_v2", "Executando apenas aplicação de cupons...")
            cupons_aplicador.aplicar_cupons(driver)
            driver.quit()
            return
        
        if args.pegar_cupons:
            log.info("run_compara_aviso_multi_v2", "Executando aplicação de cupons...")
            cupons_aplicador.aplicar_cupons(driver)
        else:
            log.info("run_compara_aviso_multi_v2", "Pulando aplicação de cupons conforme parâmetro.")

        produtos, paginas = coletar_produtos(driver, args.url)

        #print(tabulate(tabela, headers=['Produto', 'Menor Preço', 'Classificacao'], tablefmt='grid'))
        print(f"Total de páginas analisadas: {paginas}")
        print(f"Total de produtos coletados: {len(produtos)}")

        exibir_resumo_por_classificacao(produtos)
        salvar_json(produtos, nome_arquivo='mercadolivre_informatica')

        # Exemplo de caminhos
        caminho_atual = salvar_json(produtos, nome_arquivo='mercadolivre_informatica')
        caminho_anterior = os.path.join(BASE_DIR, 'ml_scraper_v2', 'ultimo.json')

        #wht_send.send_whatsapp_message(env.PHONE, env.API_KEY, f"===============")

        # Carregar e comparar
        if os.path.exists(caminho_anterior):
            json_atual = carregar_json(caminho_atual)
            json_anterior = carregar_json(caminho_anterior)
            comparar_precos_e_alertar(json_atual, json_anterior)
        else:
            print("[INFO] Nenhum arquivo anterior encontrado para comparação.")

        # Atualizar o último arquivo
        import shutil
        shutil.copy(caminho_atual, caminho_anterior)
       
        #if args.whatsapp:
        #    enviar_resumo_whatsapp(produtos)

        #caminho_xlsx = salvar_xlsx(produtos, nome_arquivo='mercadolivre_informatica')
        #enviar_para_google_drive(caminho_xlsx)
        #enviar_resumo_whatsapp_tabela(produtos)
        
        #link_google = enviar_para_google_drive(caminho_xlsx)
        #if link_google:
        #    enviar_link_whatsapp(link_google)

         # Geração do DOCX
        #caminho_docx = salvar_docx(produtos, nome_arquivo='mercadolivre_informatica')
        #link_docx = enviar_para_google_drive(caminho_docx)

        #if link_docx:
        #    enviar_link_whatsapp(link_docx)

        # Gera PDF
        caminho_pdf = salvar_pdf(produtos, nome_arquivo='mercadolivre_informatica')
        if caminho_pdf:
            link_pdf = enviar_para_google_drive(caminho_pdf)
            if link_pdf:
                enviar_link_whatsapp(link_pdf)

    finally:
        driver.quit()

if __name__ == '__main__':
    main()
